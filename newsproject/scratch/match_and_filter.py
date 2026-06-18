import docx

original_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
filtered_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"

orig_doc = docx.Document(original_path)
filt_doc = docx.Document(filtered_path)

orig_titles = []
for row in orig_doc.tables[0].rows[1:]:
    orig_titles.append(row.cells[1].text.strip())

filt_titles = []
for row in filt_doc.tables[0].rows[1:]:
    filt_titles.append(row.cells[2].text.strip())

print(f"Original titles count: {len(orig_titles)}")
print(f"Filtered titles count: {len(filt_titles)}")

print("\nFILTERED TITLES:")
for i, t in enumerate(filt_titles):
    print(f"{i+1}: {t}")

# Check which original titles are in filtered titles
matching_count = 0
for ot in orig_titles:
    # Check fuzzy match or exact match
    matched = False
    for ft in filt_titles:
        if ft.lower() in ot.lower() or ot.lower() in ft.lower():
            matched = True
            break
    if matched:
        matching_count += 1

print(f"\nOriginal titles matching filtered titles: {matching_count}")
