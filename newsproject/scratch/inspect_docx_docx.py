import docx

doc_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"
doc = docx.Document(doc_path)

print("PARAGRAPHS:")
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"P{idx}: {p.text.strip()}")

print("\nTABLES:")
print(f"Number of tables: {len(doc.tables)}")
for t_idx, table in enumerate(doc.tables):
    print(f"Table {t_idx} rows: {len(table.rows)}")
    for r_idx, row in enumerate(table.rows):
        cells = [c.text.strip() for c in row.cells]
        print(f"  Row {r_idx}: {cells}")
