import docx

doc_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"
doc = docx.Document(doc_path)

print("TOTAL PARAGRAPHS:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    print(f"P {i}: {p.text}")
