import docx
import re
import sys

sys.path.append(r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\newsproject")

original_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
filtered_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"

orig_doc = docx.Document(original_path)
filt_doc = docx.Document(filtered_path)

# Titles from the filtered table (30 articles)
filt_titles = []
for row in filt_doc.tables[0].rows[1:]:
    filt_titles.append(row.cells[2].text.strip().lower())

print(f"Filtered doc table titles count: {len(filt_titles)}")

orig_articles = []
for idx, row in enumerate(orig_doc.tables[0].rows[1:]):
    date_src = row.cells[0].text.strip()
    title = row.cells[1].text.strip()
    sector = row.cells[2].text.strip()
    direction = row.cells[3].text.strip()
    orig_articles.append((title, sector, direction, date_src))

# Let's count how many match the table exactly or fuzzily
table_matches = []
other_candidates = []

for idx, (title, sector, direction, date_src) in enumerate(orig_articles):
    matched = False
    for ft in filt_titles:
        if ft in title.lower() or title.lower() in ft:
            matched = True
            break
    if matched:
        table_matches.append((title, sector, direction, date_src))
    else:
        other_candidates.append((title, sector, direction, date_src))

print(f"Table matches count: {len(table_matches)}")
print(f"Other candidates count: {len(other_candidates)}")

# Now, let's filter other_candidates by relevance keywords
# Keep: India, Indian, TCS, Air India, seafarer, sailor, El Nino, monsoon, Sri Lanka (IT services)
# Remove: Fed, ECB, BOE, UK inflation, New Zealand, Brazil, South Africa, road accident, fire inspection, Nicobar, US tariffs on French wine, Amazon
keep_keywords = [
    r"\bIndia\b", r"\bIndian\b", r"\bTCS\b", r"\bAir India\b", r"\bseafarers\b", r"\bsailors\b",
    r"\bmonsoon\b", r"\bEl Niño\b", r"\bSri Lanka\b", r"\bBangladesh\b", r"\bJaishankar\b"
]

remove_keywords = [
    r"\bFed\b", r"\bECB\b", r"\bBOE\b", r"\bNew Zealand\b", r"\bBrazil\b", r"\bSouth Africa\b",
    r"road accident", r"fire inspection", r"Nicobar"
]

selected_others = []
for title, sector, direction, date_src in other_candidates:
    has_keep = any(re.search(pat, title, re.IGNORECASE) for pat in keep_keywords)
    has_remove = any(re.search(pat, title, re.IGNORECASE) for pat in remove_keywords)
    if has_keep and not has_remove:
        selected_others.append((title, sector, direction, date_src))

print(f"Selected others count: {len(selected_others)}")
for idx, (title, s, d, ds) in enumerate(selected_others):
    print(f"  {idx+1}: {title} | Sector: {s} | Direction: {d}")

total_kept = len(table_matches) + len(selected_others)
print(f"\nTotal Kept: {total_kept}")
