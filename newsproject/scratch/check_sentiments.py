import docx
import re
import sys

sys.path.append(r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\newsproject")

original_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
filtered_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"

orig_doc = docx.Document(original_path)
filt_doc = docx.Document(filtered_path)

filt_titles = []
for row in filt_doc.tables[0].rows[1:]:
    filt_titles.append(row.cells[2].text.strip().lower())

orig_articles = []
for idx, row in enumerate(orig_doc.tables[0].rows[1:]):
    date_src = row.cells[0].text.strip()
    title = row.cells[1].text.strip()
    sector = row.cells[2].text.strip()
    direction = row.cells[3].text.strip().lower()
    orig_articles.append((title, sector, direction, date_src))

kept = []
for title, sector, direction, date_src in orig_articles:
    matched = False
    for ft in filt_titles:
        if ft in title.lower() or title.lower() in ft:
            matched = True
            break
    
    if matched:
        kept.append((title, sector, direction, "table"))
    else:
        # Check other relevant candidates
        keep_keywords = [
            r"\bIndia\b", r"\bIndian\b", r"\bTCS\b", r"\bAir India\b", r"\bseafarers\b", r"\bsailors\b",
            r"\bmonsoon\b", r"\bEl Niño\b", r"\bSri Lanka\b", r"\bBangladesh\b", r"\bJaishankar\b"
        ]
        remove_keywords = [
            r"\bFed\b", r"\bECB\b", r"\bBOE\b", r"\bNew Zealand\b", r"\bBrazil\b", r"\bSouth Africa\b",
            r"road accident", r"fire inspection", r"Nicobar"
        ]
        has_keep = any(re.search(pat, title, re.IGNORECASE) for pat in keep_keywords)
        has_remove = any(re.search(pat, title, re.IGNORECASE) for pat in remove_keywords)
        if has_keep and not has_remove:
            kept.append((title, sector, direction, "keyword"))

pos_count = sum(1 for _, _, d, _ in kept if d == 'positive')
neg_count = sum(1 for _, _, d, _ in kept if d == 'negative')
neut_count = sum(1 for _, _, d, _ in kept if d in ('neutral', 'pending', ''))

print(f"Total Kept articles: {len(kept)}")
print(f"Positive: {pos_count}")
print(f"Negative: {neg_count}")
print(f"Neutral/Pending: {neut_count}")

# Print all articles that are kept and their status
for idx, (title, s, d, src) in enumerate(kept):
    print(f"{idx+1}: [{src.upper()}] [{d.upper()}] {title}")
