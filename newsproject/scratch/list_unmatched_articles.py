import docx
import sys

sys.path.append(r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\newsproject")

original_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
filtered_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"

orig_doc = docx.Document(original_path)
filt_doc = docx.Document(filtered_path)

filt_titles = set()
for row in filt_doc.tables[0].rows[1:]:
    filt_titles.add(row.cells[2].text.strip().lower())

print("Unmatched Articles in Original Doc:")
unmatched = []
for idx, row in enumerate(orig_doc.tables[0].rows[1:]):
    title = row.cells[1].text.strip()
    # Check if title is in filt_titles
    matched = False
    for ft in filt_titles:
        if ft in title.lower() or title.lower() in ft:
            matched = True
            break
    if not matched:
        unmatched.append((idx + 1, title, row.cells[2].text.strip()))

print(f"Total unmatched: {len(unmatched)}")
for i, t, s in unmatched:
    print(f"{i}: {t} | Sector: {s}")
