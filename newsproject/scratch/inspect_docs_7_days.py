import docx

doc_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
try:
    doc = docx.Document(doc_path)
    print("Number of tables:", len(doc.tables))
    if doc.tables:
        print("Rows in Table 0:", len(doc.tables[0].rows))
        # Print first 5 rows
        for idx, row in enumerate(doc.tables[0].rows[:10]):
            cells = [c.text.strip() for c in row.cells]
            print(f"  Row {idx}: {cells}")
except Exception as e:
    print("Error reading docs_7_Days.docx:", e)
