import docx
import re
import sqlite3
import sys

sys.path.append(r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\newsproject")

db_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\newsproject\db.sqlite3"
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

original_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\docs_7_Days.docx"
filtered_path = r"c:\Users\yadav\OneDrive\Desktop\Python\NEWS_APP\Docs_DB\Indian_Economy_News_Report_Filtered_15Jun2026.docx"

orig_doc = docx.Document(original_path)
filt_doc = docx.Document(filtered_path)

# Titles from the filtered table
filt_titles = []
for row in filt_doc.tables[0].rows[1:]:
    filt_titles.append(row.cells[2].text.strip().lower())

orig_articles = []
for idx, row in enumerate(orig_doc.tables[0].rows[1:]):
    date_src = row.cells[0].text.strip()
    title = row.cells[1].text.strip()
    orig_articles.append(title)

# Let's write the curation script logic
# Keep list
keep_titles = set()

# 1. Add all from the filtered table
for title in orig_articles:
    for ft in filt_titles:
        if ft in title.lower() or title.lower() in ft:
            keep_titles.add(title)
            break

# 2. Add other high-signal articles that are relevant
# We want exactly 52 articles kept.
# Let's list other relevant ones:
other_relevant = [
    "India's TCS partners with Anthropic to drive enterprise AI scaling - Reuters",
    "El Nino declared, all eyes now on monsoon march in India",
    "Mint Explainer | What a strong El Nino could mean for India",
    "India falls out of EM index top 10 for first time in 26 yearsWhat that means & why it matters",
    "Rubio defends US blockade after EAM Jaishankar protests seafarers deaths",
    "Be ready to respond: India on highest alert, monitoring Hormuz after 3 seafarers killed in US strike",
    "MT Marivex, Settebello, MT Jalveer: All about 3 vessels with Indian crew struck by US near Oman",
    "Failed to comply with directions: US on why it attacked 3 vessels carrying Indian crew",
    "Three Indian seafarers missing after US strike on tanker off Oman - Reuters",
    "In touch with India: US after summons to diplomat, attack on tanker off Oman",
    "India top priority for France ahead of G7 Summit; focus on West Asia and strategic defence ties",
    "COVID-19 forced vulnerable Indian households into impossible choices: Study",
    "Sri Lanka has reduced the export tariff on IT Services,",
    "Bangladesh cuts government exports tariff by 8% - effective june 2026,",
    "Wall Street IPO Boom Leaves Indias Fighting for Attention",
    "Bangladesh Targets Growth Revival With Record Budget Spending",
    "Delhi, Punjab, Odisha: States roll out free travel for NEET-UG re-exam candidates",
    "Pilots body opposes interim report on AI171 crash, demands judicial probe",
    "GE engine review delays final Air India crash report ahead of first anniversary",
    "Air India crash report delay expected due to unfinished engine analysis, source says - Reuters"
]

for t in other_relevant:
    # Find matching title in orig_articles
    for ot in orig_articles:
        if t.lower() in ot.lower() or ot.lower() in t.lower():
            keep_titles.add(ot)
            break

# Let's check how many articles we have now
print(f"Total keep titles collected: {len(keep_titles)}")

# Deduplicate "Oil prices fall" (we want to keep only the Positive one in the DB and mark the other as irrelevant)
# Let's look at the actual DB entries for kept titles and print their sentiments
kept_db_entries = []
for t in keep_titles:
    cursor.execute("SELECT id, title, direction, is_relevant FROM newsfeeds_newsarticle WHERE title = ?", (t,))
    res = cursor.fetchall()
    for r in res:
        kept_db_entries.append({
            "id": r[0],
            "title": r[1],
            "direction": r[2],
            "is_relevant": r[3]
        })

print("Kept database entries count:", len(kept_db_entries))

pos = [x for x in kept_db_entries if x['direction'] == 'positive']
neg = [x for x in kept_db_entries if x['direction'] == 'negative']
neut = [x for x in kept_db_entries if x['direction'] in ('neutral', 'pending')]

print(f"Current breakdown: Positive={len(pos)}, Negative={len(neg)}, Neutral={len(neut)}")

# Let's see if we have duplicates in kept_db_entries
seen_titles = {}
duplicates_to_remove = []
for entry in kept_db_entries:
    title = entry['title']
    if title in seen_titles:
        # We have a duplicate title! E.g. "Oil prices fall after Trump cancels strikes on Iran"
        # We want to keep the Positive one, and remove the Negative one.
        if entry['direction'] == 'negative' and seen_titles[title]['direction'] == 'positive':
            duplicates_to_remove.append(entry['id'])
        elif entry['direction'] == 'positive' and seen_titles[title]['direction'] == 'negative':
            duplicates_to_remove.append(seen_titles[title]['id'])
            seen_titles[title] = entry
    else:
        seen_titles[title] = entry

print("Duplicates to remove (IDs):", duplicates_to_remove)

# Filter out duplicates from kept_db_entries
final_kept_entries = [x for x in kept_db_entries if x['id'] not in duplicates_to_remove]
print("Final kept entries count:", len(final_kept_entries))

pos = [x for x in final_kept_entries if x['direction'] == 'positive']
neg = [x for x in final_kept_entries if x['direction'] == 'negative']
neut = [x for x in final_kept_entries if x['direction'] in ('neutral', 'pending')]

print(f"Final breakdown: Positive={len(pos)}, Negative={len(neg)}, Neutral={len(neut)}")

conn.close()
