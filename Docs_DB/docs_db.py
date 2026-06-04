import os
import sys
import sqlite3
import subprocess
from datetime import datetime, timedelta, timezone

# Auto-install python-docx if not installed
try:
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed in this environment. Installing now...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        print("python-docx successfully installed!")
    except Exception as e:
        print(f"Error: Failed to auto-install python-docx. Please install it manually with 'pip install python-docx'. Details: {e}")
        sys.exit(1)

def set_cell_background(cell, hex_color):
    """Sets the background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell internal margins (padding) in dxa (1 inch = 1440 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    for tcMar in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clear_table_borders(table):
    """Removes vertical borders and adds clean horizontal rules to the table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for tblBorders in tblPr[0].findall(qn('w:tblBorders')):
            tblPr[0].remove(tblBorders)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'bottom', 'insideH']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # border thickness (1/8 pt)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'D3D3D3')
            tblBorders.append(border)
        for border_name in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr[0].append(tblBorders)

def generate_report():
    # Setup directory and database paths
    script_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.abspath(os.path.join(script_dir, '..', 'newsproject', 'db.sqlite3'))
    output_path = os.path.join(script_dir, 'docs_7_Days.docx')

    print(f"Connecting to database at: {db_path}")
    if not os.path.exists(db_path):
        print(f"Error: Database file does not exist at '{db_path}'. Make sure paths are correct.")
        sys.exit(1)

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Calculate cutoff time (last 7 days in UTC since database is UTC)
    cutoff = datetime.now(timezone.utc).replace(tzinfo=None) - timedelta(days=7)
    cutoff_str = cutoff.strftime('%Y-%m-%d %H:%M:%S')

    # Query last 7 days data
    query = """
        SELECT title, link, source, published_date, description, summary, 
               direction, direction_reason, impact_score, event_class, 
               sector, sub_type, channel, origin
        FROM newsfeeds_newsarticle
        WHERE published_date >= ?
        ORDER BY published_date DESC
    """
    cursor.execute(query, (cutoff_str,))
    articles = cursor.fetchall()

    # Fallback to latest 7 days relative to database maximum date if current timeframe yields no data
    if not articles:
        print("No articles found in the absolute last 7 days. Checking latest data in database...")
        cursor.execute("SELECT MAX(published_date) FROM newsfeeds_newsarticle")
        max_date_str = cursor.fetchone()[0]
        if max_date_str:
            try:
                # Handle optional fractional seconds
                dt_str = max_date_str.split('.')[0]
                max_date = datetime.strptime(dt_str, '%Y-%m-%d %H:%M:%S')
                cutoff = max_date - timedelta(days=7)
                cutoff_str = cutoff.strftime('%Y-%m-%d %H:%M:%S')
                print(f"Using relative timeframe: last 7 days starting from latest DB entry: {cutoff_str}")
                cursor.execute(query, (cutoff_str,))
                articles = cursor.fetchall()
            except Exception as e:
                print(f"Failed to parse latest database date: {e}")

    if not articles:
        print("No news articles found in the database.")
        conn.close()
        return

    print(f"Fetched {len(articles)} articles. Generating document...")

    # Create document
    doc = docx.Document()

    # Set document margins (1 inch)
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(1.0)
        section.bottom_margin = docx.shared.Inches(1.0)
        section.left_margin = docx.shared.Inches(1.0)
        section.right_margin = docx.shared.Inches(1.0)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Indian Economy News Report")
    title_run.font.name = 'Georgia'
    title_run.font.size = docx.shared.Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = docx.shared.RGBColor(27, 79, 114)  # Deep Navy Blue
    title_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = docx.shared.Pt(4)

    # Document Subtitle
    subtitle_p = doc.add_paragraph()
    sub_text = (
        f"Report Timeframe: {cutoff.strftime('%d %b %Y')} to {datetime.now(timezone.utc).strftime('%d %b %Y')} (UTC)\n"
        f"Generated on: {datetime.now().strftime('%d %b %Y, %I:%M %p')} (Local Time)"
    )
    sub_run = subtitle_p.add_run(sub_text)
    sub_run.font.name = 'Calibri'
    sub_run.font.size = docx.shared.Pt(10.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = docx.shared.RGBColor(120, 120, 120)
    subtitle_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = docx.shared.Pt(18)

    # Divider line
    div_p = doc.add_paragraph()
    div_run = div_p.add_run("—" * 60)
    div_run.font.color.rgb = docx.shared.RGBColor(210, 210, 210)
    div_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    div_p.paragraph_format.space_after = docx.shared.Pt(24)

    # Executive Summary section
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.runs[0].font.name = 'Georgia'
    h1.runs[0].font.color.rgb = docx.shared.RGBColor(27, 79, 114)
    h1.runs[0].font.size = docx.shared.Pt(16)
    
    total_count = len(articles)
    positive_count = sum(1 for a in articles if str(a[6]).lower() == 'positive')
    negative_count = sum(1 for a in articles if str(a[6]).lower() == 'negative')
    neutral_count = total_count - positive_count - negative_count

    stats_p = doc.add_paragraph()
    stats_p.paragraph_format.line_spacing = 1.3
    
    run_intro = stats_p.add_run("This report summarizes parsed news events relevant to the Indian economy over the specified last 7 days. Below is a metrics breakdown of sentiments:\n\n")
    run_intro.font.name = 'Calibri'
    run_intro.font.size = docx.shared.Pt(11)

    def add_bullet_metric(p, label, value, color_rgb=None):
        r_bullet = p.add_run("   • ")
        r_bullet.font.name = 'Calibri'
        r_bullet.font.size = docx.shared.Pt(11)
        
        r_label = p.add_run(label)
        r_label.font.name = 'Calibri'
        r_label.font.bold = True
        r_label.font.size = docx.shared.Pt(11)
        
        r_val = p.add_run(f": {value}\n")
        r_val.font.name = 'Calibri'
        r_val.font.size = docx.shared.Pt(11)
        if color_rgb:
            r_label.font.color.rgb = color_rgb
            r_val.font.color.rgb = color_rgb

    add_bullet_metric(stats_p, "Total Articles Analyzed", total_count)
    add_bullet_metric(stats_p, "Positive Impact Sentiments", positive_count, docx.shared.RGBColor(30, 132, 73))
    add_bullet_metric(stats_p, "Negative Impact Sentiments", negative_count, docx.shared.RGBColor(203, 67, 53))
    add_bullet_metric(stats_p, "Neutral/Pending Sentiments", neutral_count, docx.shared.RGBColor(100, 110, 120))

    # Add spacing before table
    doc.add_paragraph().paragraph_format.space_after = docx.shared.Pt(6)

    # Overview Table section
    h2 = doc.add_heading("2. Articles Overview", level=1)
    h2.runs[0].font.name = 'Georgia'
    h2.runs[0].font.color.rgb = docx.shared.RGBColor(27, 79, 114)
    h2.runs[0].font.size = docx.shared.Pt(16)
    
    table = doc.add_table(rows=1, cols=4)
    clear_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    headers = ["Date & Source", "Article Title", "Sector & Class", "Direction"]
    col_widths = [docx.shared.Inches(1.5), docx.shared.Inches(3.2), docx.shared.Inches(1.8), docx.shared.Inches(1.0)]

    # Style Header Row
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1B4F72")  # Brand color
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        hdr_cells[i].width = col_widths[i]
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = docx.shared.Pt(10)
        run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)

    # Populate table rows
    for idx, art in enumerate(articles):
        row_cells = table.add_row().cells
        
        # Format Date & Source
        date_raw = art[3]
        try:
            dt_parsed = datetime.strptime(date_raw.split('.')[0], '%Y-%m-%d %H:%M:%S')
            date_display = dt_parsed.strftime('%d %b %Y')
        except Exception:
            date_display = date_raw
        source = art[2]
        
        # Populate Date & Source
        row_cells[0].text = f"{date_display}\n({source})"
        
        # Populate Title
        row_cells[1].text = art[0]
        
        # Populate Sector & Event Class
        sector = (art[10] or "General").replace('_', ' ')
        ev_class = (art[9] or "Economic Impact").replace('_', ' ')
        row_cells[2].text = f"{sector}\n[{ev_class}]"
        
        # Populate Direction
        direction = (art[6] or "neutral").capitalize()
        row_cells[3].text = direction

        # Apply styles to cells
        for col_idx, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.width = col_widths[col_idx]
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = docx.shared.Pt(9.5)
            
            # Shading alternate rows (zebra striping)
            if idx % 2 == 1:
                set_cell_background(cell, "F9EBEA" if direction.lower() == 'negative' and col_idx == 3 else ("EAFAF1" if direction.lower() == 'positive' and col_idx == 3 else "F4F6F6"))
            elif direction.lower() == 'negative' and col_idx == 3:
                set_cell_background(cell, "FDEDEC")
            elif direction.lower() == 'positive' and col_idx == 3:
                set_cell_background(cell, "E8F8F5")
                
        # Direction styling
        dir_p = row_cells[3].paragraphs[0]
        dir_run = dir_p.runs[0]
        dir_run.font.bold = True
        if direction.lower() == 'positive':
            dir_run.font.color.rgb = docx.shared.RGBColor(30, 132, 73)
        elif direction.lower() == 'negative':
            dir_run.font.color.rgb = docx.shared.RGBColor(203, 67, 53)
        else:
            dir_run.font.color.rgb = docx.shared.RGBColor(120, 120, 120)

    # Spacing
    doc.add_paragraph().paragraph_format.space_after = docx.shared.Pt(18)

    # Detailed Analysis Section
    h3 = doc.add_heading("3. Detailed News Insights", level=1)
    h3.runs[0].font.name = 'Georgia'
    h3.runs[0].font.color.rgb = docx.shared.RGBColor(27, 79, 114)
    h3.runs[0].font.size = docx.shared.Pt(16)

    for idx, art in enumerate(articles):
        title, link, source, published_date, description, summary, direction, direction_reason, impact_score, event_class, sector, sub_type, channel, origin = art
        
        # Article Header (Heading 2)
        art_heading = doc.add_heading(f"3.{idx+1} {title}", level=2)
        art_heading.runs[0].font.name = 'Georgia'
        art_heading.runs[0].font.size = docx.shared.Pt(13.5)
        art_heading.runs[0].font.color.rgb = docx.shared.RGBColor(33, 97, 140)
        art_heading.paragraph_format.space_before = docx.shared.Pt(16)
        art_heading.paragraph_format.space_after = docx.shared.Pt(4)

        # Meta attributes
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.line_spacing = 1.2
        meta_p.paragraph_format.space_after = docx.shared.Pt(8)
        
        r_meta = meta_p.add_run(
            f"Source: {source} | Published: {published_date} | Origin: {origin or 'Global'}\n"
            f"Sector: {sector.replace('_', ' ')} | Event Class: {event_class.replace('_', ' ')}\n"
            f"Direction: {direction.capitalize()}"
        )
        r_meta.font.name = 'Calibri'
        r_meta.font.size = docx.shared.Pt(9)
        r_meta.font.italic = True
        r_meta.font.color.rgb = docx.shared.RGBColor(110, 110, 110)

        # Insight heading
        ins_h = doc.add_heading("Insight & Summary", level=3)
        ins_h.runs[0].font.name = 'Georgia'
        ins_h.runs[0].font.size = docx.shared.Pt(11)
        ins_h.runs[0].font.color.rgb = docx.shared.RGBColor(50, 50, 50)
        ins_h.paragraph_format.space_before = docx.shared.Pt(6)
        ins_h.paragraph_format.space_after = docx.shared.Pt(2)

        # Insight paragraph
        ins_p = doc.add_paragraph()
        ins_p.paragraph_format.left_indent = docx.shared.Inches(0.2)
        ins_p.paragraph_format.space_after = docx.shared.Pt(6)
        r_ins = ins_p.add_run(summary or "No summary/insight available.")
        r_ins.font.name = 'Calibri'
        r_ins.font.size = docx.shared.Pt(10.5)

        # Reason heading
        if direction_reason and direction_reason.strip() != '':
            reason_h = doc.add_heading("Sentiment Justification", level=3)
            reason_h.runs[0].font.name = 'Georgia'
            reason_h.runs[0].font.size = docx.shared.Pt(11)
            reason_h.runs[0].font.color.rgb = docx.shared.RGBColor(50, 50, 50)
            reason_h.paragraph_format.space_before = docx.shared.Pt(6)
            reason_h.paragraph_format.space_after = docx.shared.Pt(2)

            # Reason paragraph
            reason_p = doc.add_paragraph()
            reason_p.paragraph_format.left_indent = docx.shared.Inches(0.2)
            reason_p.paragraph_format.space_after = docx.shared.Pt(6)
            r_reason = reason_p.add_run(direction_reason)
            r_reason.font.name = 'Calibri'
            r_reason.font.size = docx.shared.Pt(10.5)

        # Link (if applicable)
        if link and link != '#':
            link_p = doc.add_paragraph()
            link_p.paragraph_format.space_after = docx.shared.Pt(12)
            r_link_lbl = link_p.add_run("Original Source Link: ")
            r_link_lbl.font.name = 'Calibri'
            r_link_lbl.font.size = docx.shared.Pt(9.5)
            r_link_lbl.font.bold = True
            
            r_link_val = link_p.add_run(link)
            r_link_val.font.name = 'Calibri'
            r_link_val.font.size = docx.shared.Pt(9.5)
            r_link_val.font.color.rgb = docx.shared.RGBColor(41, 128, 185)

    # Save to file, overwriting existing
    doc.save(output_path)
    print(f"Success! Generated report written and replaced old data at: {output_path}")

    conn.close()

if __name__ == "__main__":
    generate_report()
